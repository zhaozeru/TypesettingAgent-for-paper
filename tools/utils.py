from langchain_community.chat_models import ChatZhipuAI
from config.setting import ZHIPUAI_API_KEY
import docx
import re
import json
from typing import Dict, Any

def str_to_bool(s: str) -> bool:
    s = s.strip().lower()
    if s == "true":
        return True
    elif s == "false":
        return False
    else:
        raise ValueError(f"Invalid boolean string: {s}")


def zhipu_llm(
    model: str = "glm-4.5-air",
    api_key: str = ZHIPUAI_API_KEY,
    thinking: str = "enabled",
    stream: bool = True,
    temperature: float = 0.7,
    timeout: float = 30.0,
    **kwargs,
):
    return ChatZhipuAI(
        api_key=api_key,
        model=model,
        thinking={
            "type": thinking,
        },
        stream=stream,
        temperature=temperature,
        timeout=timeout,
        **kwargs
    )


def read_docx(file_path: str) -> list:
    """
    读取 .docx 文件，返回带索引的段落列表
    每一项格式：{"index": 0, "text": "原文", "style": "Normal", "is_heading": False, "heading_level": None}
    """
    try:
        doc = docx.Document(file_path)
        lines = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 保留引号，避免干扰 JSON
            text = text.replace('"', '&quot;')

            # 判断是否为标题
            style = para.style.name if para.style else "Normal"
            is_heading = style.lower().startswith("heading") or style in ["标题 1", "标题 2", "标题 3"]
            heading_level = None
            if is_heading:
                if "1" in style or "一" in style:
                    heading_level = 1
                elif "2" in style or "二" in style:
                    heading_level = 2
                elif "3" in style or "三" in style:
                    heading_level = 3

            lines.append({
                "index": i,
                "text": text,
                "style": style,
                "is_heading": is_heading,
                "heading_level": heading_level
            })

        # 表格处理（可选：也可单独作为结构）
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    lines.append({
                        "index": len(lines),
                        "text": f"[表格] {row_text}",
                        "style": "Table",
                        "is_heading": False,
                        "heading_level": None
                    })

        return lines

    except Exception as e:
        raise RuntimeError(f"解析 .docx 文件失败: {str(e)}")


def str_to_dict(text):
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return {}


def parse_to_json(text: str):
    if not text or not text.strip():
        return {}
    try:
        data = json.loads(text.strip())
        return data
    except json.JSONDecodeError:
        pass

    code_match = re.search(r"```(?:json|JSON)?\s*([\s\S]+?)\s*```", text, re.DOTALL)
    if code_match:
        json_text = code_match.group(1).strip()
    else:
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1:
            start = text.find("[")
            end = text.rfind("]")
        if start == -1 or end == -1:
            return {}
        json_text = text[start : end + 1].strip()

    try:
        json_text = json_text.replace('&', '\\"')
        data = json.loads(json_text)
        return data
    except json.JSONDecodeError as e:
        return {}

def clean_dict(d):
    if not isinstance(d, dict):
        return print("❌ 字典清洗函数处理对象不为字典格式！")
    cleaned = {}
    for key, value in d.items():
        if value is None and key in ["agent_mission", "agent_result"]:
            continue
        if isinstance(value, list):
            new_list = []
            for item in value:
                if isinstance(item, dict):
                    content = item.get("content")
                    request = item.get("request")
                    if content and request and len(request) > 0:
                        new_list.append(item)
            if new_list:
                cleaned[key] = new_list
        elif isinstance(value, dict):
            cleaned_sub = clean_dict(value)
            if cleaned_sub:
                cleaned[key] = cleaned_sub
        else:
            cleaned[key] = value

    return cleaned if cleaned else None

def read_input_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        print("✅ 文件读取成功！")
        return content
    except FileNotFoundError:
        print("❌ 错误：文件未找到，请检查路径是否正确。")
    except PermissionError:
        print("❌ 错误：没有权限读取该文件。")
    except Exception as e:
        print(f"❌ 读取失败：{e}")
    return None

import json
from typing import Any, Dict, List
from datetime import datetime, date
from decimal import Decimal
from pathlib import Path
import enum

# LangChain 相关
from langchain_core.messages import BaseMessage
from langchain_core.documents import Document  # 如果状态中包含 Document


class MessageJSONEncoder(json.JSONEncoder):
    """
    专门用于序列化 LangGraph 状态的 JSON Encoder。
    处理 BaseMessage、datetime、set、Path、Enum、Decimal 等非标准类型。
    """

    def default(self, obj: Any) -> Any:
        # --- 1. 处理 BaseMessage 及其子类 ---
        if isinstance(obj, BaseMessage):
            # 提取所有可用信息
            message_data: Dict[str, Any] = {
                "type": obj.__class__.__name__,  # "AIMessage", "HumanMessage", "ToolMessage"
                "content": obj.content,
                "additional_kwargs": obj.additional_kwargs,
                "response_metadata": obj.response_metadata,
                # 👇 常用扩展字段
                "id": getattr(obj, "id", None),  # 消息 ID
                "name": getattr(obj, "name", None),  # 发送者名称
                "example": getattr(obj, "example", False),  # 是否是示例
            }

            # 👇 特定类型的消息添加额外字段
            if hasattr(obj, "tool_calls") and obj.tool_calls:
                message_data["tool_calls"] = obj.tool_calls
            if hasattr(obj, "invalid_tool_calls") and obj.invalid_tool_calls:
                message_data["invalid_tool_calls"] = obj.invalid_tool_calls
            if hasattr(obj, "usage_metadata") and obj.usage_metadata:
                message_data["usage_metadata"] = obj.usage_metadata

            return message_data

        # --- 2. 处理 Document ---
        if isinstance(obj, Document):
            return {
                "page_content": obj.page_content,
                "metadata": obj.metadata,
                "type": "Document"
            }

        # --- 3. 处理 datetime 和 date ---
        if isinstance(obj, (datetime, date)):
            return obj.isoformat()  # 转成标准字符串 "2025-08-15T09:30:00"

        # --- 4. 处理 set ---
        if isinstance(obj, set):
            return list(obj)  # 转成 list

        # --- 5. 处理 Path ---
        if isinstance(obj, Path):
            return str(obj)  # 转成字符串路径

        # --- 6. 处理 Enum ---
        if isinstance(obj, enum.Enum):
            return obj.value  # 返回枚举值

        # --- 7. 处理 Decimal ---
        if isinstance(obj, Decimal):
            return float(obj)  # 或 str(obj) 保留精度

        # --- 8. 处理 bytes ---
        if isinstance(obj, bytes):
            try:
                return obj.decode('utf-8')  # 尝试转成字符串
            except UnicodeDecodeError:
                return obj.hex()  # 转成十六进制字符串

        # --- 9. 处理其他未知对象 ---
        # 如果以上都不是，尝试返回其 __dict__ 或 str 表示
        try:
            # 尝试获取对象的字典表示
            if hasattr(obj, '__dict__'):
                return obj.__dict__
            # 或者返回字符串表示
            return repr(obj)
        except Exception:
            # 最后兜底
            return f"<{type(obj).__name__}>"

        # 如果都不行，让父类处理（通常会抛 TypeError）
        return super().default(obj)
#
# if __name__ == "__main__":
#     try:
#         # 创建模型实例
#         llm = zhipu_llm()
#
#         # 调用模型（非流式测试更简单，先关掉 stream）
#         result = llm.invoke("请用中文介绍一下你自己")
#         print("✅ 调用成功！")
#         print(result.content)
#
#     except Exception as e:
#         print("❌ 出错了：", str(e))
