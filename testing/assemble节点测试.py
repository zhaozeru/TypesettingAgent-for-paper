from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
import os
from typing import Dict, List, Any
state1 = {
    "agent_mission": [
    {
        "section": "首页",
        "subsection": "论文题目",
        "content": "基于《金匮要略》的多粒度知识表示模型构建研究",
        "para_index": 0,
        "request": [
            {
                "字体": "宋体"
            },
            {
                "字号": "三号"
            },
            {
                "加粗": "是"
            },
            {
                "对齐": "居中"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "作者姓名",
        "content": "王世文  赵泽儒*  邢建斌",
        "para_index": 1,
        "request": [
            {
                "字体": "宋体"
            },
            {
                "字号": "小四"
            },
            {
                "对齐": "居中"
            },
            {
                "行距": "1.5倍行距"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "作者单位",
        "content": "（天津师范大学 管理学院，天津  300387）",
        "para_index": 2,
        "request": [
            {
                "字体": "宋体"
            },
            {
                "字号": "五号"
            },
            {
                "对齐": "居中"
            },
            {
                "行距": "1.5倍行距"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "摘要",
        "content": "摘要：[目的] 本文基于多粒度层级划分提出多粒度的知识表示框架与模型，以提升中医古籍知识挖掘与利用效率。[方法] 本文首先通过对《金匮要略》的分析与归纳，提出从知识内容角度出发的多粒度层级划分结构。然后，据此构建出中医古籍多粒度知识表示框架，并给出多粒度知识表示模型，并以《金匮要略》和《温病条辨》中的知识内容进行了模型展示与验证。[结果] 对多粒度知识表示模型的应用，展示出模型的普适性和科学性，体现出模型在中医古籍领域中的可用性和有效性，有效支持中医古籍知识表示与组织。[结论] 多粒度知识表示模型为中医古籍知识挖掘与组织提供了可行方案，具有理论与实践价值。",
        "para_index": 3,
        "request": [
            {
                "字体": "宋体"
            },
            {
                "字号": "小四"
            },
            {
                "行距": "1.5倍行距"
            },
            {
                "格式": "【目的/意义】……【方法/过程】……【结果/结论】……【创新/局限】"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "关键词",
        "content": "关键词：中医古籍；知识表示；多粒度；知识元；知识表示模型",
        "para_index": 4,
        "request": [
            {
                "字体": "宋体"
            },
            {
                "字号": "五号"
            },
            {
                "行距": "1.5倍"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "中图分类号",
        "content": "分类号：G254",
        "para_index": 5,
        "request": [
            {
                "字体": "宋体"
            },
            {
                "字号": "五号"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "英文题目",
        "content": "Construction of a Multi-Granularity Knowledge Representation Model Based on Synopsis of Golden Chamber",
        "para_index": 7,
        "request": [
            {
                "字体": "Times New Roman"
            },
            {
                "字号": "三号"
            },
            {
                "加粗": "是"
            },
            {
                "对齐": "居中"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "英文作者姓名",
        "content": "WANG Shi-wen  ZHAO Ze-ru*",
        "para_index": 8,
        "request": [
            {
                "字体": "Times New Roman"
            },
            {
                "字号": "小四"
            },
            {
                "对齐": "居中"
            },
            {
                "行距": "1.5倍行距"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "英文作者单位",
        "content": "(College of Management‌, Tianjin Normal University, Tianjin, 300387)",
        "para_index": 9,
        "request": [
            {
                "字体": "Times New Roman"
            },
            {
                "字号": "五号"
            },
            {
                "对齐": "居中"
            },
            {
                "行距": "1.5倍行距"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "英文摘要",
        "content": "Abstract：[Objective] This paper proposes a multi-granularity knowledge representation framework and model based on multi-granularity hierarchical division to enhance the efficiency of knowledge mining and utilization in ancient Chinese medical literature.[Methods] This article first proposes a multi granularity hierarchical structure from the perspective of knowledge content by analyzing and summarizing the \"quot;Synopsis of the Golden Chamber\"quot;. Then, based on this, a multi granularity knowledge representation framework for traditional Chinese medicine ancient books was constructed, and a multi granularity knowledge representation model was proposed. The model was demonstrated and validated using the knowledge content from \"quot;Synopsis of Golden Chamber\"quot; and \"quot;Wenbing Tiaobian\"quot;.[Results] The application of multi granularity knowledge representation models demonstrates the universality and scientificity of the models, reflecting their usability and effectiveness in the field of traditional Chinese medicine ancient books, and effectively supporting the representation and organization of traditional Chinese medicine ancient book knowledge.[Conclusions] The multi-granularity knowledge representation model provides a feasible solution for knowledge mining and organization in ancient Chinese medical literature, with both theoretical and practical significance.",
        "para_index": 10,
        "request": [
            {
                "字体": "Times New Roman"
            },
            {
                "字号": "小四"
            },
            {
                "行距": "1.5倍行距"
            },
            {
                "格式": "【目的/意义】……【方法/过程】……【结果/结论】……【创新/局限】"
            }
        ]
    },
    {
        "section": "首页",
        "subsection": "英文关键词",
        "content": "Keywords： TCM ancient literature; knowledge representation; multi-granularity; knowledge element; Description model of knowledge",
        "para_index": 11,
        "request": [
            {
                "字体": "Times New Roman"
            },
            {
                "字号": "五号"
            },
            {
                "行距": "1.5倍"
            }
        ]
    },
    {
        "section": "其他排版要求",
        "subsection": "未在文档中出现的部分",
        "request": [
            {
                "subsection": "页面布局",
                "missing_request": []
            },
            {
                "subsection": "副题目",
                "missing_request": [
                    {
                        "对齐": "右对齐"
                    }
                ]
            },
            {
                "subsection": "正文段落",
                "missing_request": [
                    {
                        "字体": "宋体"
                    },
                    {
                        "字号": "小四"
                    },
                    {
                        "行距": "1.5倍"
                    },
                    {
                        "段落格式": "首行缩进2字符"
                    }
                ]
            },
            {
                "subsection": "一级标题",
                "missing_request": [
                    {
                        "字体": "黑体"
                    },
                    {
                        "字号": "三号"
                    },
                    {
                        "加粗": "是"
                    },
                    {
                        "对齐": "居中"
                    },
                    {
                        "段前段后": "各12pt"
                    }
                ]
            },
            {
                "subsection": "二级标题",
                "missing_request": [
                    {
                        "字体": "黑体"
                    },
                    {
                        "字号": "四号"
                    },
                    {
                        "加粗": "是"
                    },
                    {
                        "对齐": "左对齐"
                    },
                    {
                        "段前段后": "各6pt"
                    }
                ]
            },
            {
                "subsection": "三级标题",
                "missing_request": [
                    {
                        "字体": "楷体"
                    },
                    {
                        "字号": "小四"
                    },
                    {
                        "加粗": "是"
                    },
                    {
                        "对齐": "左对齐"
                    }
                ]
            },
            {
                "subsection": "图表",
                "missing_request": []
            },
            {
                "subsection": "公式",
                "missing_request": []
            },
            {
                "subsection": "算法",
                "missing_request": []
            },
            {
                "subsection": "代码片段",
                "missing_request": []
            },
            {
                "subsection": "文献列表",
                "missing_request": []
            },
            {
                "subsection": "作者简介",
                "missing_request": []
            },
            {
                "subsection": "作者贡献声明",
                "missing_request": []
            },
            {
                "subsection": "课题或基金项目",
                "missing_request": []
            }
        ]
    }
]
}
state2 ={
        "agent_result": {
            "status": "success",
            "layout_result": [
                {
                    "section": "首页",
                    "subsection": "论文题目",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "id": "call_-8395464655852354728",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 16
                            },
                            "id": "call_-8395464655852354727",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "id": "call_-8395464655852354726",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "id": "call_-8395464655852354725",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395464655852354728",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "宋体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395464655852354727",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 16
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 16.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395464655852354726",
                            "tool_name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "bold": True
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395464655852354725",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "center"
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "作者姓名",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "id": "call_-8395447716500134966",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "id": "call_-8395447716500134965",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "id": "call_-8395447716500134964",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "id": "call_-8395447716500134963",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395447716500134966",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "宋体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395447716500134965",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 12.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395447716500134964",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "center"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395447716500134963",
                            "tool_name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "line_spacing": 1.5
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "作者单位",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "id": "call_-8395467404632013971",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "id": "call_-8395467404632013970",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "id": "call_-8395467404632013969",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "id": "call_-8395467404632013968",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395467404632013971",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "宋体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467404632013970",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 10.5
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467404632013969",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "center"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467404632013968",
                            "tool_name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "line_spacing": 1.5
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "摘要",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "id": "call_-8395455206923517540",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "id": "call_-8395455206923517539",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "id": "call_-8395455206923517538",
                            "type": "tool_call"
                        },
                        {
                            "name": "check_abstract_structure",
                            "args": {
                                "abstract": "摘要：[目的] 本文基于多粒度层级划分提出多粒度的知识表示框架与模型，以提升中医古籍知识挖掘与利用效率。[方法] 本文首先通过对《金匮要略》的分析与归纳，提出从知识内容角度出发的多粒度层级划分结构。然后，据此构建出中医古籍多粒度知识表示框架，并给出多粒度知识表示模型，并以《金匮要略》和《温病条辨》中的知识内容进行了模型展示与验证。[结果] 对多粒度知识表示模型的应用，展示出模型的普适性和科学性，体现出模型在中医古籍领域中的可用性和有效性，有效支持中医古籍知识表示与组织。[结论] 多粒度知识表示模型为中医古籍知识挖掘与组织提供了可行方案，具有理论与实践价值。",
                                "required_sections": [
                                    "目的/意义",
                                    "方法/过程",
                                    "结果/结论",
                                    "创新/局限"
                                ]
                            },
                            "id": "call_-8395455206923517537",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395455206923517540",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "宋体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395455206923517539",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 12.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395455206923517538",
                            "tool_name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "line_spacing": 1.5
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395455206923517537",
                            "tool_name": "check_abstract_structure",
                            "args": {
                                "abstract": "摘要：[目的] 本文基于多粒度层级划分提出多粒度的知识表示框架与模型，以提升中医古籍知识挖掘与利用效率。[方法] 本文首先通过对《金匮要略》的分析与归纳，提出从知识内容角度出发的多粒度层级划分结构。然后，据此构建出中医古籍多粒度知识表示框架，并给出多粒度知识表示模型，并以《金匮要略》和《温病条辨》中的知识内容进行了模型展示与验证。[结果] 对多粒度知识表示模型的应用，展示出模型的普适性和科学性，体现出模型在中医古籍领域中的可用性和有效性，有效支持中医古籍知识表示与组织。[结论] 多粒度知识表示模型为中医古籍知识挖掘与组织提供了可行方案，具有理论与实践价值。",
                                "required_sections": [
                                    "目的/意义",
                                    "方法/过程",
                                    "结果/结论",
                                    "创新/局限"
                                ]
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "failed",
                                    "type": "check",
                                    "data": {
                                        "result": False,
                                        "check_name": "摘要结构检查",
                                        "message": "摘要缺少以下内容部分：创新/局限",
                                        "severity": "error",
                                        "details": {
                                            "found_sections": [
                                                "目的/意义",
                                                "方法/过程",
                                                "结果/结论"
                                            ],
                                            "missing_sections": [
                                                "创新/局限"
                                            ]
                                        },
                                        "revision_comment": {
                                            "author": "内容检查器",
                                            "text": "摘要内容不完整，缺少：创新/局限 部分",
                                            "paragraph_index": None,
                                            "run_index": None
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "关键词",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "id": "call_-8395465033809567243",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "id": "call_-8395465033809567242",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "id": "call_-8395465033809567241",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395465033809567243",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "宋体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395465033809567242",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 10.5
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395465033809567241",
                            "tool_name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "line_spacing": 1.5
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "中图分类号",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "id": "call_-8395465033809575606",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "id": "call_-8395465033809575605",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395465033809575606",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "宋体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395465033809575605",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 10.5
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "英文题目",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "id": "call_-8395438233211846510",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 16
                            },
                            "id": "call_-8395438233211846509",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "id": "call_-8395438233211846508",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "id": "call_-8395438233211846507",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395438233211846510",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "Times New Roman"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395438233211846509",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 16
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 16.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395438233211846508",
                            "tool_name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "bold": True
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395438233211846507",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "center"
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "英文作者姓名",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "id": "call_-8395468023107607470",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "id": "call_-8395468023107607469",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "id": "call_-8395468023107607468",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "id": "call_-8395468023107607467",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395468023107607470",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "Times New Roman"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395468023107607469",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 12.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395468023107607468",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "center"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395468023107607467",
                            "tool_name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "line_spacing": 1.5
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "英文作者单位",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "id": "call_-8395467885668580020",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "id": "call_-8395467885668580019",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "id": "call_-8395467885668580018",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "id": "call_-8395467885668580017",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395467885668580020",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "Times New Roman"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668580019",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 10.5
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668580018",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "center"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668580017",
                            "tool_name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "line_spacing": 1.5
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "英文摘要",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "id": "call_-8395468469784362369",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "id": "call_-8395468469784362368",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "id": "call_-8395468469784362367",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395468469784362369",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "Times New Roman"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395468469784362368",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 12.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395468469784362367",
                            "tool_name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "line_spacing": 1.5
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "首页",
                    "subsection": "英文关键词",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "id": "call_-8395447716500134903",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "id": "call_-8395447716500134902",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "id": "call_-8395447716500134901",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395447716500134903",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "Times New Roman"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "Times New Roman"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395447716500134902",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 10.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 10.5
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395447716500134901",
                            "tool_name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "line_spacing": 1.5
                                        }
                                    }
                                }
                            }
                        }
                    ]
                },
                {
                    "section": "其他排版要求",
                    "subsection": "未在文档中出现的部分",
                    "status": "success",
                    "raw_model_response": "",
                    "tool_calls": [
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "right"
                            },
                            "id": "call_-8395467885668571329",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "id": "call_-8395467885668571328",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "id": "call_-8395467885668571327",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "id": "call_-8395467885668571326",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_indent",
                            "args": {
                                "first_line": 2
                            },
                            "id": "call_-8395467885668571325",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "黑体"
                            },
                            "id": "call_-8395467885668571324",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 16
                            },
                            "id": "call_-8395467885668571323",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "id": "call_-8395467885668571322",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "id": "call_-8395467885668571321",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "黑体"
                            },
                            "id": "call_-8395467885668571320",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 14
                            },
                            "id": "call_-8395467885668571319",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "id": "call_-8395467885668571318",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "left"
                            },
                            "id": "call_-8395467885668571317",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_name",
                            "args": {
                                "font_name": "楷体"
                            },
                            "id": "call_-8395467885668571316",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "id": "call_-8395467885668571315",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "id": "call_-8395467885668571314",
                            "type": "tool_call"
                        },
                        {
                            "name": "set_paragraph_align",
                            "args": {
                                "align": "left"
                            },
                            "id": "call_-8395467885668571313",
                            "type": "tool_call"
                        }
                    ],
                    "tool_execution_results": [
                        {
                            "tool_call_id": "call_-8395467885668571329",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "right"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "right"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571328",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "宋体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "宋体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571327",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 12.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571326",
                            "tool_name": "set_line_spacing",
                            "args": {
                                "multiple": 1.5
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "line_spacing": 1.5
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571325",
                            "tool_name": "set_paragraph_indent",
                            "args": {
                                "first_line": 2
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "indent": {
                                                "first_line": 2.0
                                            }
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571324",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "黑体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "黑体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571323",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 16
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 16.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571322",
                            "tool_name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "bold": True
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571321",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "center"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "center"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571320",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "黑体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "黑体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571319",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 14
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 14.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571318",
                            "tool_name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "bold": True
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571317",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "left"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "left"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571316",
                            "tool_name": "set_font_name",
                            "args": {
                                "font_name": "楷体"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_name": "楷体"
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571315",
                            "tool_name": "set_font_size",
                            "args": {
                                "pt": 12
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "font_size": 12.0
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571314",
                            "tool_name": "set_font_bold",
                            "args": {
                                "bold": True
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_font",
                                    "data": {
                                        "styles": {
                                            "bold": True
                                        }
                                    }
                                }
                            }
                        },
                        {
                            "tool_call_id": "call_-8395467885668571313",
                            "tool_name": "set_paragraph_align",
                            "args": {
                                "align": "left"
                            },
                            "execution_result": {
                                "status": "success",
                                "result": {
                                    "status": "success",
                                    "action": "set_paragraph",
                                    "data": {
                                        "styles": {
                                            "align": "left"
                                        }
                                    }
                                }
                            }
                        }
                    ]
                }
            ]
        }
}
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# -----------------------------
# 配置常量
# -----------------------------

# 中文字号转磅值 (pt)
CHINESE_TO_PTS = {
    "初号": 42,
    "一号": 26,
    "二号": 22,
    "三号": 16,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
}

# 对齐方式映射
ALIGN_MAP = {
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
}

# 所有执行型工具（格式设置）
EXECUTION_TOOLS = {
    "set_font_name",
    "set_font_size",
    "set_font_bold",
    "set_font_italic",
    "set_font_underline",
    "set_paragraph_align",
    "set_line_spacing",
    "set_space_before",
    "set_space_after",
    "set_paragraph_indent",  # ✅ 新增：段落缩进
}

# 所有检查型工具
CHECK_TOOLS = {
    "check_title_length",
    "check_keyword_count",
    "check_abstract_structure",
    "check_abstract_length",
    "check_author_affiliation_by_llm",
    "check_author_name_format",
    "check_paragraph_structure",
}

# 输出路径
OUTPUT_DIR = r"D:\my_document\pycharm\project\ai_agent_project\src\05paper_layout_agent\experiment\output"
OUTPUT_FILE = "cover_document.docx"


# -----------------------------
# 工具函数
# -----------------------------

def get_lookup_key(item):
    """生成唯一键"""
    section = item.get("section", "").strip()
    subsection = item.get("subsection", "").strip()
    if section and subsection:
        return f"{section}_{subsection}"
    return section or subsection


def resolve_font_size(pt):
    """解析字号：支持字符串（如'小四'）或数字"""
    if isinstance(pt, str):
        return CHINESE_TO_PTS.get(pt, 12)
    return float(pt)


# -----------------------------
# 主函数
# -----------------------------

def generate_document(state1, state2):
    agent_mission = state1.get("agent_mission", [])
    layout_result = state2.get("agent_result", {}).get("layout_result", [])

    # 1. 构建内容映射
    content_map = {}
    for item in agent_mission:
        if "content" not in item:
            print(f"⚠️ 跳过：缺少 'content' -> {item.get('section', '')}/{item.get('subsection', '')}")
            continue
        key = get_lookup_key(item)
        if not key:
            print(f"⚠️ 跳过：无法确定键 -> {item}")
            continue
        content_map[key] = item["content"]
        print(f"📝 映射内容: {key} -> {item['content'][:50]}...")

    # 2. 创建文档
    doc = Document()
    revision_comments = []

    # 3. 遍历 layout_result 应用格式
    processed_keys = set()

    for item in layout_result:
        key = get_lookup_key(item)
        if not key:
            print(f"⚠️ layout_result 项缺少键: {item}")
            continue

        content = content_map.get(key)
        if not content:
            print(f"⚠️ 未找到内容（key: {key}），跳过...")
            continue

        p = doc.add_paragraph()
        run = p.add_run(content)
        applied_count = 0
        check_results = []

        for exec_res in item.get("tool_execution_results", []):
            tool_name = exec_res.get("tool_name")
            if not tool_name:
                continue

            args = exec_res.get("args", {})
            result = exec_res.get("execution_result", {}).get("result", {})
            status = result.get("status")

            try:
                # --- 检查型工具 ---
                if tool_name in CHECK_TOOLS:
                    data = result.get("data", {})
                    comment = data.get("revision_comment")
                    check_info = {
                        "tool": tool_name,
                        "status": status,
                        "message": data.get("message", "无信息"),
                        "result": data.get("result", False),
                        "severity": data.get("severity", "info"),
                        "comment": comment
                    }
                    check_results.append(check_info)
                    if comment:
                        revision_comments.append({
                            "section": item.get("section", ""),
                            "subsection": item.get("subsection", ""),
                            "comment": comment.get("text", "无注释")
                        })
                    print(f"🔍 检查 [{tool_name}]: {check_info['message']} ({status})")

                # --- 执行型工具 ---
                elif tool_name in EXECUTION_TOOLS and status == "success":
                    if tool_name == "set_font_name":
                        font_name = args.get("font_name")
                        if font_name:
                            run.font.name = font_name
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                            applied_count += 1

                    elif tool_name == "set_font_size":
                        pt_val = args.get("pt")
                        size_pt = resolve_font_size(pt_val)
                        run.font.size = Pt(size_pt)
                        applied_count += 1

                    elif tool_name == "set_font_bold":
                        run.bold = bool(args.get("bold", False))
                        applied_count += 1

                    elif tool_name == "set_font_italic":
                        run.italic = bool(args.get("italic", False))
                        applied_count += 1

                    elif tool_name == "set_font_underline":
                        run.underline = bool(args.get("underline", False))
                        applied_count += 1

                    elif tool_name == "set_paragraph_align":
                        align = args.get("align")
                        if align in ALIGN_MAP:
                            p.alignment = ALIGN_MAP[align]
                            applied_count += 1

                    elif tool_name == "set_line_spacing":
                        p.paragraph_format.line_spacing = float(args.get("multiple", 1.0))
                        applied_count += 1

                    elif tool_name == "set_space_before":
                        space = args.get("pt", 0)
                        p.paragraph_format.space_before = Pt(float(space))
                        applied_count += 1

                    elif tool_name == "set_space_after":
                        space = args.get("pt", 0)
                        p.paragraph_format.space_after = Pt(float(space))
                        applied_count += 1
                    elif tool_name == "set_paragraph_indent":
                        # 支持首行缩进（characters）或悬挂缩进
                        indent_type = args.get("type", "first_line")  # "first_line" or "hanging"
                        characters = args.get("characters", 0)  # 字符数
                        pt_per_char = 10.5 / 2  # 估算：小四字宽 ≈ 0.5字符 ≈ 5.25pt
                        indent_pt = characters * pt_per_char

                        if indent_type == "first_line":
                            p.paragraph_format.first_line_indent = Pt(indent_pt)
                        elif indent_type == "hanging":
                            p.paragraph_format.left = Pt(indent_pt)
                            p.paragraph_format.first_line_indent = Pt(-indent_pt)
                        applied_count += 1

                else:
                    if tool_name in EXECUTION_TOOLS:
                        print(f"⚠️ 执行失败: {tool_name}")

            except Exception as e:
                print(f"❌ 应用工具 '{tool_name}' 失败: {e}")

        print(f"✅ 添加 '{key}'，应用 {applied_count} 项格式，完成 {len(check_results)} 项检查")
        processed_keys.add(key)

    # 4. 检查未处理的内容
    for key in content_map:
        if key not in processed_keys:
            print(f"⚠️ 注意：内容 '{key}' 未在 layout_result 中处理")

    # 5. 添加修订建议到文档
    if revision_comments:
        doc.add_paragraph()
        heading = doc.add_paragraph("📌 修订建议")
        heading.style = "Heading 1"
        for rc in revision_comments:
            p = doc.add_paragraph()
            p.add_run(f"{rc['subsection']}: ").bold = True
            p.add_run(rc["comment"])

        print(f"\n📋 已将 {len(revision_comments)} 条修订建议插入文档末尾")

    # 6. 保存文档
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    filepath = os.path.join(OUTPUT_DIR, OUTPUT_FILE)

    try:
        doc.save(filepath)
        abs_path = os.path.abspath(filepath)
        print(f"\n🎉 Word 文档已成功生成：{abs_path}")
        print(f"📄 文档包含 {len(doc.paragraphs)} 个段落")

        # 统计
        total_applied = sum(
            1 for item in layout_result
            for exec in item.get("tool_execution_results", [])
            if exec.get("tool_name") in EXECUTION_TOOLS
            and exec.get("execution_result", {}).get("result", {}).get("status") == "success"
        )
        total_checks = sum(
            1 for item in layout_result
            for exec in item.get("tool_execution_results", [])
            if exec.get("tool_name") in CHECK_TOOLS
        )
        print(f"📊 总计: {total_applied} 项格式应用, {total_checks} 项检查完成")

    except PermissionError:
        print(f"❌ 保存失败：文件可能被占用，请关闭 '{OUTPUT_FILE}' 后重试。")
    except Exception as e:
        print(f"❌ 保存失败：{str(e)}")

    return doc

generate_document(state1, state2)