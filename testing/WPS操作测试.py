from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐

# 分段设置格式
run1 = p.add_run("今天是7月30日，")
run1.bold = True
run1.italic = True

run2 = p.add_run("风和日丽")
run2.bold = True
run2.italic = True
run2.font.color.rgb = RGBColor(0, 255, 0)  # 绿色

run3 = p.add_run("的一天！")
run3.bold = True
run3.italic = True

doc.save("text.docx")
print("✅ 文档已生成：text.docx")
