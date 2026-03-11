from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.shared import Inches
import os
import random
from typing import Dict, List, Any
from langchain_core.messages import AIMessage, SystemMessage

# 中文字号映射表
CHINESE_TO_PTS = {
    "初号": 42, "一号": 26, "二号": 22, "三号": 16,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
}

# 对齐方式映射
ALIGN_MAP = {
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
}

# 文档组装工具集
DOCUMENT_TOOLS = {
    # 字体格式工具
    "set_font_name", "set_font_size", "set_font_bold",
    "set_font_italic", "set_font_underline",

    # 段落格式工具
    "set_paragraph_align", "set_line_spacing",
    "set_space_before", "set_space_after", "set_paragraph_indent",

    # 文档结构工具
    "add_section_break", "set_page_margins", "add_page_number"
}


def resolve_font_size(size: str) -> float:
    """将中文或字符串转为 Pt 值"""
    if isinstance(size, (int, float)):
        return float(size)
    return CHINESE_TO_PTS.get(size, 12)  # 默认小四


def apply_document_tool(tool_name: str, args: dict, doc, p, run):
    """应用文档工具到文档、段落或运行文本"""
    try:
        if tool_name == "set_font_name":
            font_name = args.get("font_name")
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        elif tool_name == "set_font_size":
            pt_val = args.get("pt")
            size_pt = resolve_font_size(pt_val)
            run.font.size = Pt(size_pt)

        elif tool_name == "set_font_bold":
            run.bold = bool(args.get("bold", False))

        elif tool_name == "set_font_italic":
            run.italic = bool(args.get("italic", False))

        elif tool_name == "set_font_underline":
            run.underline = bool(args.get("underline", False))

        elif tool_name == "set_paragraph_align":
            align = args.get("align")
            if align in ALIGN_MAP:
                p.alignment = ALIGN_MAP[align]

        elif tool_name == "set_line_spacing":
            p.paragraph_format.line_spacing = float(args.get("multiple", 1.0))

        elif tool_name == "set_space_before":
            p.paragraph_format.space_before = Pt(float(args.get("pt", 0)))

        elif tool_name == "set_space_after":
            p.paragraph_format.space_after = Pt(float(args.get("pt", 0)))

        elif tool_name == "set_paragraph_indent":
            indent_type = args.get("type", "first_line")
            characters = args.get("characters", 0)
            pt_per_char = 10.5 / 2
            indent_pt = characters * pt_per_char

            if indent_type == "first_line":
                p.paragraph_format.first_line_indent = Pt(indent_pt)
            elif indent_type == "hanging":
                p.paragraph_format.left = Pt(indent_pt)
                p.paragraph_format.first_line_indent = Pt(-indent_pt)

        elif tool_name == "add_section_break":
            doc.add_section(WD_SECTION.NEW_PAGE)

        elif tool_name == "set_page_margins":
            section = doc.sections[-1]
            for margin in ["top", "bottom", "left", "right"]:
                if margin in args:
                    setattr(section, f"{margin}_margin", Inches(float(args[margin])))

        elif tool_name == "add_page_number":
            section = doc.sections[-1]
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run(f"第 {args.get('page', '')} 页")

    except Exception as e:
        raise RuntimeError(f"工具 {tool_name} 应用失败: {e}")


def add_issues_section(doc, verification_issues):
    """在文档末尾添加问题汇总部分"""
    if not verification_issues:
        return

    # 添加分节符
    doc.add_section(WD_SECTION.NEW_PAGE)

    # 添加标题
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("待改进问题汇总")
    run_title.font.size = Pt(16)
    run_title.bold = True
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加说明
    p_desc = doc.add_paragraph()
    p_desc.add_run("以下是在文档检查过程中发现的需要改进的问题：")

    # 添加问题列表
    for i, issue in enumerate(verification_issues, 1):
        p_issue = doc.add_paragraph()
        p_issue.paragraph_format.left_indent = Pt(20)
        p_issue.paragraph_format.space_after = Pt(6)

        run_issue = p_issue.add_run(f"{i}. 【{issue.get('section', '未知')}-{issue.get('subsection', '未知')}】")
        run_issue.bold = True

        p_desc2 = doc.add_paragraph()
        p_desc2.paragraph_format.left_indent = Pt(40)
        p_desc2.add_run(f"问题类型: {issue.get('issue_type', '未知')}")

        p_desc3 = doc.add_paragraph()
        p_desc3.paragraph_format.left_indent = Pt(40)
        p_desc3.add_run(f"问题描述: {issue.get('description', '无描述')}")

        p_desc4 = doc.add_paragraph()
        p_desc4.paragraph_format.left_indent = Pt(40)
        run_suggestion = p_desc4.add_run(f"改进建议: {issue.get('suggestion', '无建议')}")
        run_suggestion.font.color.rgb = RGBColor(0, 128, 0)  # 绿色


def assemble_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """文档组装车间 - 将各个部件组装成完整文档"""
    system_messages = [
        SystemMessage(content="<-- 进入 09 文档组装车间 -->")
    ]

    # 获取封面数据和验证问题
    cover = state.get("cover", {})
    verification_issues = state.get("verification_issues", [])

    agent_mission = cover.get("agent_mission", [])
    layout_result = cover.get("agent_result", {}).get("layout_result", [])

    if not agent_mission:
        system_messages.append(SystemMessage(content="⚠️ 无排版任务，生成空文档"))
        # 仍然创建文档，但内容为空

    # 创建 Word 文档
    doc = Document()

    # 设置默认页面边距
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # 构建内容映射
    content_map = {}
    for item in agent_mission:
        content = item.get("content")
        if content:
            key = item.get("subsection") or item.get("section")
            if key:
                content_map[key] = content

    # 组装文档主体内容
    assembly_log = []
    for item in layout_result:
        subsection = item.get("subsection")
        section_name = item.get("section")
        lookup_key = subsection or section_name

        content = content_map.get(lookup_key)
        if not content:
            continue

        # 添加新段落
        p = doc.add_paragraph()
        run = p.add_run(content)

        # 应用所有格式工具
        tool_count = 0
        for exec_res in item.get("tool_execution_results", []):
            tool_name = exec_res.get("tool_name")
            if tool_name in DOCUMENT_TOOLS:
                try:
                    apply_document_tool(
                        tool_name,
                        exec_res.get("args", {}),
                        doc, p, run
                    )
                    tool_count += 1
                except Exception as e:
                    assembly_log.append(f"❌ {lookup_key} - {tool_name}: {e}")

        assembly_log.append(f"✅ {lookup_key}: 应用 {tool_count} 个工具")

    # 添加验证问题汇总
    if verification_issues:
        add_issues_section(doc, verification_issues)
        assembly_log.append(f"📝 添加了 {len(verification_issues)} 个待改进问题")

    # 保存文档
    output_dir = r"D:\Project\agent_project\src\05paper_layout_agent\experiment\output"
    os.makedirs(output_dir, exist_ok=True)

    random_suffix = random.randint(1000, 9999)
    filename = f"文档组装结果_{random_suffix}.docx"
    output_file = os.path.join(output_dir, filename)

    try:
        doc.save(output_file)
        abs_path = os.path.abspath(output_file)

        # 添加组装日志到系统消息
        for log in assembly_log:
            system_messages.append(SystemMessage(content=log))

        system_messages.append(SystemMessage(content=f"🎉 文档组装完成: {abs_path}"))

        return {
            "messages": system_messages,
            "final_paper": abs_path,
            "awaiting": "user_email",
            "assembly_log": assembly_log,
            "issues_count": len(verification_issues)
        }

    except Exception as e:
        system_messages.append(SystemMessage(content=f"❌ 文档保存失败: {str(e)}"))
        return {
            "messages": system_messages,
            "error": f"文档保存失败: {str(e)}",
            "awaiting": ""
        }